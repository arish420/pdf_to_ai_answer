import re
import pdfplumber
import pytesseract
from pdf2image import convert_from_path
from docx import Document
import openai
import docx
import os
import tempfile
import streamlit as st
import config
import time
from dotenv import load_dotenv

# Load environment variables
load_dotenv() 

# Function to extract text from a scanned PDF (OCR)
def extract_text_with_ocr(pdf_path):
    try:
        images = convert_from_path(pdf_path)
        full_text = ""
        for image in images:
            full_text += pytesseract.image_to_string(image) + "\n"
        return full_text
    except Exception as e:
        st.error(f"Error processing PDF with OCR: {str(e)}")
        return ""

# Function to extract text from a selectable text PDF
def extract_text_from_pdf(pdf_path):
    text = ""
    try:
        with pdfplumber.open(pdf_path) as pdf:
            for page in pdf.pages:
                extracted_text = page.extract_text()
                if extracted_text:
                    text += extracted_text + "\n"
    except Exception as e:
        st.error(f"Error extracting text from PDF: {str(e)}")
    return text

# Function to determine if the PDF is scanned or text-based
def get_pdf_text(pdf_path):
    text = extract_text_from_pdf(pdf_path)
    if len(text.strip()) < 50:  # If text is too small, assume it's a scanned PDF
        st.info("PDF appears to be scanned. Using OCR...")
        return extract_text_with_ocr(pdf_path)
    else:
        st.info("PDF has selectable text. Extracting directly...")
        return text

# Function to extract text from Word document
def extract_text_from_docx(docx_path):
    try:
        doc = Document(docx_path)
        text = "" 
        for para in doc.paragraphs:
            text += para.text + "\n"
        return text
    except Exception as e:
        st.error(f"Error extracting text from DOCX: {str(e)}")
        return ""

# Function to extract questions from the text using regex
def extract_questions(text):
    # This pattern ensures that the question is more than a single word
    question_pattern = r"([A-Z][^?.!]*\?)"
    questions = re.findall(question_pattern, text)

    # Filter out single word questions or short questions
    filtered_questions = [q.strip() for q in questions if len(q.split()) > 2]  # Ensure at least 3 words
    return filtered_questions

# Main function to extract questions from PDF
def extract_questions_from_pdf(pdf_path):
    full_text = get_pdf_text(pdf_path)
    return extract_questions(full_text)

# Main function to extract questions from DOCX
def extract_questions_from_docx(docx_path):
    full_text = extract_text_from_docx(docx_path)
    return extract_questions(full_text)

# Update the OpenAI API key configuration
def generate_answers(questions, api_key=None):
    # Use config API key if none provided
    openai.api_key = api_key or config.OPENAI_API_KEY
    qa_pairs = []
    
    progress_bar = st.progress(0)
    status_text = st.empty()
    
    for i, question in enumerate(questions):
        status_text.text(f"Generating answer for question {i+1} of {len(questions)}")
        try:
            response = openai.ChatCompletion.create(
                model="gpt-4",
                messages=[
                    {"role": "system", "content": "You are a domain expert providing factual, concise answers. Never mention AI, LLMs, or language models in your responses. Never say 'As an AI' or similar phrases. Respond in a natural, human-like manner with factual information only."},
                    {"role": "user", "content": question}
                ],
                max_tokens=150
            )
            answer = response.choices[0].message.content.strip()
            qa_pairs.append((question, answer))
        except Exception as e:
            st.error(f"Error generating answer for question: {str(e)}")
            qa_pairs.append((question, "Unable to generate an answer at this time."))
        
        # Update progress
        progress_bar.progress((i + 1) / len(questions))
    
    status_text.empty()
    progress_bar.empty()
    return qa_pairs

# Function to save Q&A to a DOCX file
def save_qa_to_docx(questions_answers, output_file):
    doc = Document()
    doc.add_heading("Questions & Answers from Document", level=1)
    
    for i, (question, answer) in enumerate(questions_answers, start=1):
        doc.add_paragraph(f"Question {i}: {question}", style="Heading 2")
        doc.add_paragraph(f"{answer}\n")
    
    doc.save(output_file)
    return output_file

# Add local path handling
def ensure_temp_dir():
    temp_dir = os.path.join(os.path.dirname(__file__), 'temp')
    if not os.path.exists(temp_dir):
        os.makedirs(temp_dir)
    return temp_dir

def get_temp_file_path(suffix):
    temp_dir = ensure_temp_dir()
    return os.path.join(temp_dir, f"temp_{str(int(time.time()))}{suffix}")
